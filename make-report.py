#!/usr/bin/env python3
"""Generate a stylized Give A Bit accounting report as PDF, Excel, and Word."""
import json, datetime, os, sys
sys.path.insert(0, "/root/hq")
from report_data import main as fetch_data

LOGO = "/root/hq/brand-logo.png"
INK = (0x0E/255, 0x1B/255, 0x2A/255)      # #0E1B2A
BRASS = (0xB8/255, 0x89/255, 0x3A/255)    # #B8893A
PAPER = (0xF4/255, 0xEF/255, 0xE4/255)    # #F4EFE4
GREEN = (0x22/255, 0xC5/255, 0x5E/255)

def build_pdf(data, path):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.pdfgen import canvas
    from reportlab.lib.utils import ImageReader

    W, H = A4
    c = canvas.Canvas(path, pagesize=A4)
    c.setFillColorRGB(*PAPER); c.rect(0,0,W,H,fill=1,stroke=0)

    # Header band
    c.setFillColorRGB(*INK); c.rect(0, H-38*mm, W, 38*mm, fill=1, stroke=0)
    # Logo
    try:
        img = ImageReader(LOGO); iw, ih = img.getSize()
        scale = 30*mm / iw
        c.drawImage(img, 15*mm, H-33*mm, width=iw*scale, height=ih*scale, mask='auto')
    except Exception as e:
        pass
    c.setFillColorRGB(1,1,1); c.setFont("Helvetica-Bold", 16)
    c.drawString(60*mm, H-18*mm, "Give A Bit — Accounting Report")
    c.setFont("Helvetica", 9); c.setFillColorRGB(0.85,0.85,0.85)
    c.drawString(60*mm, H-23*mm, "Executive financial summary across the family of projects")

    # Generated date
    c.setFillColorRGB(*INK); c.setFont("Helvetica", 9)
    c.drawString(15*mm, H-45*mm, f"Generated: {data['generated']}")

    y = H - 55*mm
    def section(title):
        nonlocal y
        c.setFillColorRGB(*BRASS); c.setFont("Helvetica-Bold", 11)
        c.drawString(15*mm, y, title); y -= 6*mm
    def kv(label, value):
        nonlocal y
        c.setFillColorRGB(*INK); c.setFont("Helvetica", 10)
        c.drawString(18*mm, y, label)
        c.setFont("Helvetica-Bold", 10)
        c.drawRightString(W-15*mm, y, str(value)); y -= 5.5*mm
    def rule():
        nonlocal y
        c.setStrokeColorRGB(*BRASS); c.setLineWidth(0.6)
        c.line(15*mm, y, W-15*mm, y); y -= 4*mm

    section("Key Figures")
    kv("Total .ots stamps", data['stamps_total'])
    kv("Confirmed stamps", data['stamps_confirmed'])
    kv("Payments recorded", data['payments_total'])
    kv("Total sats received (IN)", f"{data['total_sats_in']:,}")
    rule()

    section("Stamps by Client")
    for name, cnt in sorted(data['clients'].items(), key=lambda x:-x[1]):
        if name: kv(name or "(none)", cnt)
    rule()

    section("Pricing")
    kv("Base price per timestamp (sats)", data['base_price'])
    kv("Rate adjustment (%)", data.get('rate_adj', 0))
    kv("Effective price per timestamp (sats)", data.get('eff_price', data['base_price']))

    c.setFillColorRGB(*GREEN); c.setFont("Helvetica", 7)
    c.drawString(15*mm, 12*mm, "Give A Bit · self-custody, no-KYC, sovereign · Confidential")
    c.showPage(); c.save()
    return path

def build_excel(data, path):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    wb = Workbook(); ws = wb.active; ws.title = "Report"
    ws.sheet_view.showGridLines = False
    ws.column_dimensions['A'].width = 34; ws.column_dimensions['B'].width = 22
    FILL = PatternFill("solid", fgColor="0E1B2A")
    ws["A1"] = "Give A Bit — Accounting Report"; ws["A1"].font = Font(bold=True, size=14, color="0E1B2A")
    ws["A2"] = f"Generated {data['generated']}"; ws["A2"].font = Font(italic=True, color="6B7280", size=9)
    ws["A4"] = "Key Figures"; ws["A4"].fill = FILL; ws["A4"].font = Font(bold=True, color="FFFFFF")
    fig = [
        ("Total .ots stamps", data['stamps_total']),
        ("Confirmed", data['stamps_confirmed']),
        ("Payments", data['payments_total']),
        ("Total sats IN", data['total_sats_in']),
        ("Base price/stamp", data['base_price']),
    ]
    for i,(k,v) in enumerate(fig,5):
        ws.cell(row=i,column=1,value=k)
        ws.cell(row=i,column=2,value=v)
    ws["A11"] = "Stamps by Client"; ws["A11"].fill = FILL; ws["A11"].font = Font(bold=True, color="FFFFFF")
    r=12
    for name,cnt in sorted(data['clients'].items(), key=lambda x:-x[1]):
        if name:
            ws.cell(row=r,column=1,value=name); ws.cell(row=r,column=2,value=cnt); r+=1
    wb.save(path); return path

def build_word(data, path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    doc = Document()
    try:
        doc.add_picture(LOGO, width=Inches(2.0))
    except: pass
    doc.add_heading("Give A Bit — Accounting Report", 0)
    doc.add_paragraph(f"Generated {data['generated']}")
    h = doc.add_heading("Key Figures", level=1)
    for k in ["Total .ots stamps","Confirmed stamps","Payments recorded","Total sats IN","Base price/stamp"]:
        p = doc.add_paragraph(); p.add_run(f"{k}: ").bold = True
        v = data['stamps_total'] if "stamps" in k and "Confirmed" not in k and "price" not in k else \
            data['stamps_confirmed'] if "Confirmed" in k else \
            data['payments_total'] if "Payments" in k else \
            data['base_price'] if "price" in k else data['total_sats_in']
        p.add_run(str(v))
    doc.add_heading("Stamps by Client", level=1)
    for name,cnt in sorted(data['clients'].items(), key=lambda x:-x[1]):
        if name:
            doc.add_paragraph(f"{name}: {cnt}", style="List Bullet")
    doc.save(path); return path

if __name__ == "__main__":
    data = fetch_data()
    outdir = "/root/giveabit-reports"
    os.makedirs(outdir, exist_ok=True)
    stamp = datetime.datetime.now().strftime("%Y%m%d-%H%M")
    pdf = build_pdf(data, f"{outdir}/giveabit-report-{stamp}.pdf")
    xls = build_excel(data, f"{outdir}/giveabit-report-{stamp}.xlsx")
    doc = build_word(data, f"{outdir}/giveabit-report-{stamp}.docx")
    print("PDF:", pdf)
    print("XLSX:", xls)
    print("DOCX:", doc)
