#!/usr/bin/env python3
"""Give A Bit — branded accounting report generator (master + per-site).

Produces stylized PDF / Excel / Word reports with:
  - strong identification (org, site, report type, report ID)
  - generation timestamp + report period (scope)
  - scope section (what's covered, data sources, method, caveats)
  - brand logo, navy header, Give A Bit palette
  - live data from Google Sheets ledgers
"""
import json, datetime, os, sys, re
sys.path.insert(0, "/root/hq")
from report_data import get_creds, fetch_sheet

from googleapiclient.discovery import build

LOGO = "/root/hq/brand-logo.png"
MASTER = "14wV3zYz9OgJ0K9mWohq3mRGZxRWjlEj4FGUpTl1C1iA"
INK = (0x0E/255, 0x1B/255, 0x2A/255)
BRASS = (0xB8/255, 0x89/255, 0x3A/255)
PAPER = (0xF4/255, 0xEF/255, 0xE4/255)
GREEN = (0x22/255, 0xC5/255, 0x5E/255)
GREY = (0x6B/255, 0x72/255, 0x80/255)

SITE_INFO = {
    "motopass":    {"folder": "1tqgZtmFPdpXi959AU3YmTyh9hicQiXBu", "desc": "Passport & travel-document service", "domain": "motopass.giveabit.io"},
    "satohash":    {"folder": "1Cl_6ZGZ36mE4pWWkC1yKwQCvNPrGfe__", "desc": "OpenTimestamps / Bitcoin proof plane", "domain": "satohash.io"},
    "katoa":       {"folder": "1c4oNe_wsRIcgygJVojvhT05WsCjfTGUe", "desc": "Community / giving project", "domain": "katoa.org"},
    "tadbuy":      {"folder": "1BHa_Wr_Pg3OMt8SowJ3nQc55vs3euJXu", "desc": "Commerce / marketplace", "domain": "tadbuy.giveabit.io"},
    "sherpacarta": {"folder": "1Ah-Wou0-AzSvhli37gbotE33XmnXR5p7", "desc": "Digital rights product guide", "domain": "sherpacarta.org"},
    "stranded":    {"folder": "1qJi3-YIbY4fpd_qo_ELEfjyyKzWfVcCL", "desc": "Community / giving project", "domain": "stranded.giveabit.io"},
}

def iso_date(s):
    if not s: return None
    m = re.match(r"(\d{4})-(\d{2})-(\d{2})", str(s))
    if m: return datetime.date(*[int(x) for x in m.groups()])
    return None

def build_scope_text(period_start=None, period_end=None, extra=None):
    lines = [
        f"Scope: This report covers {period_start or 'all recorded'} to {period_end or 'present'}.",
        "Data sources: Give A Bit Google Sheets ledgers (MoneyFlow per-site + MASTER-LEDGER) fed by the Satohash stamp API.",
        "Method: Live OTS stamp counts and ledger balances pulled programmatically; totals are as-recorded, not audited.",
        "Caveats: Currently the family operates on a free tier (REQUIRE_LIGHTNING=false) — 'sats received' reflect real Lightning movements only. Historical stamps prior to the accounting system (2026-08-21) may be under-counted in the money ledgers.",
    ]
    if extra: lines.append(extra)
    return lines

def gather_master_data(sht):
    ledger = fetch_sheet(sht, MASTER, 'Ledger!A2:J1000')
    stamps = [r for r in ledger if r and r[0]]
    stamps_confirmed = sum(1 for r in stamps if len(r)>3 and r[3]=="confirmed")
    from collections import Counter
    clients = Counter((r[7] if len(r)>7 else "?") for r in stamps)
    pays = fetch_sheet(sht, MASTER, 'Payments!A2:L1000')
    pays = [r for r in pays if r and r[0] and str(r[0])[:4].isdigit()]
    total_sats_in = sum(int(r[5]) for r in pays if len(r)>5 and str(r[5]).replace(",","").isdigit())
    # rate control
    acct = fetch_sheet(sht, MASTER, 'Accounting!B5:B10')
    try:
        base = int(float(acct[0][0])); rate = int(float(acct[1][0])); eff = int(float(acct[2][0]))
        btc_usd = float(acct[3][0]) if len(acct)>3 and acct[3] and str(acct[3][0]).replace('.','').isdigit() else 77863
        sats_per_btc = int(float(acct[4][0])) if len(acct)>4 and acct[4] else 100000000
    except Exception:
        base, rate, eff = 5, 0, 5
        btc_usd, sats_per_btc = 77863, 100000000
    return {
        "stamps_total": len(stamps),
        "stamps_confirmed": stamps_confirmed,
        "clients": dict(clients),
        "payments_total": len(pays),
        "total_sats_in": total_sats_in,
        "base_price": base, "rate_adj": rate, "eff_price": eff,
        "btc_usd": btc_usd, "sats_per_btc": sats_per_btc,
    }

def gather_site_data(sht, site):
    """Pull a site's stamps (from master ledger by client) + its MoneyFlow."""
    ledger = fetch_sheet(sht, MASTER, 'Ledger!A2:J1000')
    # identify site stamps by client/filename
    site_stamps = []
    for r in ledger:
        if not r or not r[0]: continue
        cl = (r[7] if len(r)>7 else "").lower()
        fn = (r[2] if len(r)>2 else "").lower()
        if site in cl or site in fn:
            site_stamps.append(r)
    # MoneyFlow via the site's spreadsheet (IMPORTRANGE data we can also read directly)
    # We'll read master Payments filtered by site (col B = site)
    pays = fetch_sheet(sht, MASTER, 'Payments!A2:L1000')
    site_pays = [r for r in pays if r and r[0] and str(r[0])[:4].isdigit() and (len(r)>1 and (r[1] or "").lower()==site)]
    site_sats_in = sum(int(r[5]) for r in site_pays if len(r)>5 and str(r[5]).replace(",","").isdigit())
    return {
        "stamps_total": len(site_stamps),
        "stamps_confirmed": sum(1 for r in site_stamps if len(r)>3 and r[3]=="confirmed"),
        "payments": len(site_pays),
        "sats_in": site_sats_in,
    }

def gen_report_id(report_type, site=None):
    ts = datetime.datetime.now().strftime("%Y%m%d")
    n = datetime.datetime.now().strftime("%H%M")
    base = "GAB"
    scope = site.upper() if site else "MASTER"
    return f"{base}-{scope}-{report_type.upper()}-{ts}-{n}"

def stamp_now():
    return datetime.datetime.now().strftime("%Y-%m-%d %H:%M")

# ── PDF ──────────────────────────────────────────────────────────────
def build_pdf(title, site_name, report_id, generated, scope_lines, sections, path, trend=None):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.pdfgen import canvas
    from reportlab.lib.utils import ImageReader
    from reportlab.lib.colors import HexColor
    W, H = A4
    INK_C = HexColor("#0E1B2A"); BRASS_C = HexColor("#B8893A")
    PAPER_C = HexColor("#F4EFE4"); GREEN_C = HexColor("#22C55E")
    GREY_C = HexColor("#6B7280")
    c = canvas.Canvas(path, pagesize=A4)
    c.setFillColor(PAPER_C); c.rect(0,0,W,H,fill=1,stroke=0)

    # ── Header band ──
    c.setFillColor(INK_C); c.rect(0, H-42*mm, W, 42*mm, fill=1, stroke=0)
    try:
        img = ImageReader(LOGO); iw, ih = img.getSize()
        sc = 32*mm/iw
        c.drawImage(img, 15*mm, H-36*mm, width=iw*sc, height=ih*sc, mask='auto')
    except Exception: pass
    c.setFillColorRGB(1,1,1); c.setFont("Helvetica-Bold", 19)
    c.drawString(58*mm, H-20*mm, title[:50])
    c.setFont("Helvetica", 10); c.setFillColorRGB(0.88,0.88,0.88)
    c.drawString(58*mm, H-25*mm, "Give A Bit · Sovereign accounting & reporting")
    # identification block
    c.setFont("Helvetica", 9)
    c.drawString(58*mm, H-31*mm, f"Report ID:  {report_id}")
    c.drawString(58*mm, H-35.5*mm, f"Generated:  {generated}")

    # ── Report title line ──
    c.setFillColor(INK_C); c.setFont("Helvetica-Bold", 12)
    c.drawString(15*mm, H-50*mm, f"Report:  {site_name}")

    y = H - 58*mm
    def section(title, color=BRASS_C):
        nonlocal y
        if y < 42*mm:
            c.showPage(); c.setFillColor(PAPER_C); c.rect(0,0,W,H,fill=1,stroke=0); y = H-20*mm
        # brass underline bar
        c.setFillColor(color); c.rect(15*mm, y-1*mm, W-30*mm, 2*mm, fill=1, stroke=0)
        c.setFillColor(INK_C); c.setFont("Helvetica-Bold", 12.5)
        c.drawString(15*mm, y-4.5*mm, title)
        y -= 11*mm

    def bullet(label, text_lines, indent=24, size=9.5, step=6.8, label_step=6.8):
        """Render a labeled bullet: bold brass label on its own line, then
        wrapped body text below with generous line spacing (12pt for 9.5pt font)."""
        nonlocal y
        # label on its own line
        if y < 45*mm:
            c.showPage(); c.setFillColor(PAPER_C); c.rect(0,0,W,H,fill=1,stroke=0); y = H-20*mm
        c.setFillColor(BRASS_C); c.setFont("Helvetica-Bold", size)
        c.drawString(15*mm, y, label)
        y -= label_step
        # wrapped body text
        c.setFillColor(INK_C); c.setFont("Helvetica", size)
        maxw = W - (indent+3)*mm
        words = str(text_lines).split()
        cur = ""
        for w in words:
            if c.stringWidth((cur+" "+w).strip(), "Helvetica", size) < maxw:
                cur = (cur+" "+w).strip()
            else:
                if y < 45*mm:
                    c.showPage(); c.setFillColor(PAPER_C); c.rect(0,0,W,H,fill=1,stroke=0); y = H-20*mm
                c.drawString(indent*mm, y, cur); y -= step
                cur = w
        if cur:
            if y < 45*mm:
                c.showPage(); c.setFillColor(PAPER_C); c.rect(0,0,W,H,fill=1,stroke=0); y = H-20*mm
            c.drawString(indent*mm, y, cur); y -= step
        y -= 3*mm

    def kv(label, value, size=10):
        nonlocal y
        if y < 42*mm:
            c.showPage(); c.setFillColor(PAPER_C); c.rect(0,0,W,H,fill=1,stroke=0); y = H-20*mm
        c.setFillColor(GREY_C); c.setFont("Helvetica", size)
        c.drawString(20*mm, y, label)
        c.setFillColor(INK_C); c.setFont("Helvetica-Bold", size)
        c.drawRightString(W-20*mm, y, str(value)); y -= 6*mm

    def rule():
        nonlocal y
        if y < 42*mm:
            c.showPage(); c.setFillColor(PAPER_C); c.rect(0,0,W,H,fill=1,stroke=0); y = H-20*mm
        c.setStrokeColor(HexColor("#D8D2C4")); c.setLineWidth(0.6)
        c.line(15*mm, y, W-15*mm, y); y -= 6*mm

    # ── Report Scope (labeled bullets) ──
    section("Report Scope")
    # scope_lines currently are "Label: text" — split them
    for line in scope_lines:
        if ":" in line:
            lab, rest = line.split(":", 1)
            bullet(lab.strip() + ":", rest.strip())
        else:
            bullet("•", line)
    rule()

    def table(header, rows):
        """Render a simple column table with header row."""
        nonlocal y
        n = len(header)
        left = 15*mm
        total_w = W - 30*mm
        # proportional widths: first col wider
        widths = [total_w*0.30] + [total_w*0.70/(n-1) for _ in range(n-1)] if n > 1 else [total_w]
        row_h = 5.5*mm
        # header
        if y - row_h < 42*mm:
            c.showPage(); c.setFillColor(PAPER_C); c.rect(0,0,W,H,fill=1,stroke=0); y = H-20*mm
        c.setFillColor(INK_C); c.setFont("Helvetica-Bold", 8)
        x = left
        for i, htxt in enumerate(header):
            c.drawString(x+1*mm, y, str(htxt)[:18]); x += widths[i]
        y -= row_h
        # data rows
        for row in rows:
            if y - row_h < 42*mm:
                c.showPage(); c.setFillColor(PAPER_C); c.rect(0,0,W,H,fill=1,stroke=0); y = H-20*mm
            x = left
            for i, cell in enumerate(row):
                c.setFillColor(INK_C); c.setFont("Helvetica", 8)
                txt = str(cell)
                # right-align numeric-ish cells (not first col)
                if i > 0 and txt.replace(",","").replace(".","").isdigit():
                    c.drawRightString(x+widths[i]-1*mm, y, txt)
                else:
                    c.drawString(x+1*mm, y, txt[:18])
                x += widths[i]
            # subtle row separator
            c.setStrokeColor(HexColor("#E5DFD2")); c.setLineWidth(0.4)
            c.line(left, y-1*mm, W-15*mm, y-1*mm)
            y -= row_h
        y -= 3*mm

    for title2, rows in sections:
        section(title2)
        if isinstance(rows, list) and rows and isinstance(rows[0], (list, tuple)):
            # If first row is all strings and >2 cols, treat as a table (header+data)
            if len(rows[0]) > 2 and all(isinstance(c, str) for c in rows[0]):
                header = rows[0]
                data = rows[1:]
                table(header, data)
            else:
                for row in rows:
                    if len(row)==2:
                        kv(str(row[0]), str(row[1]))
                    else:
                        bullet("•", str(row[0]))
        else:
            for it in rows:
                bullet("•", str(it))
        y -= 2*mm

    # ── Trend chart (monthly stamps bar chart) ──
    if trend and trend.get("monthly"):
        monthly = trend["monthly"]
        if len(monthly) >= 1:
            section("Stamp Volume Trend")
            months_l = sorted(monthly.keys())
            vals = [monthly[m] for m in months_l]
            maxv = max(vals) or 1
            chart_left = 25*mm; chart_bottom = y - 38*mm
            chart_w = W - 50*mm; chart_h = 34*mm
            # axis
            c.setStrokeColor(HexColor("#0E1B2A")); c.setLineWidth(1)
            c.line(chart_left, chart_bottom, chart_left+chart_w, chart_bottom)
            c.line(chart_left, chart_bottom, chart_left, chart_bottom+chart_h)
            n = len(vals)
            bar_w = min(12*mm, (chart_w / n) * 0.6)
            gap = chart_w / n
            for i, (m, v) in enumerate(zip(months_l, vals)):
                bh = (v / maxv) * chart_h
                bx = chart_left + gap*i + (gap-bar_w)/2
                c.setFillColor(HexColor("#B8893A"))
                c.rect(bx, chart_bottom, bar_w, bh, fill=1, stroke=0)
                # value label
                c.setFillColor(HexColor("#0E1B2A")); c.setFont("Helvetica-Bold", 7)
                c.drawCentredString(bx+bar_w/2, chart_bottom+bh+1*mm, str(v))
                # month label
                c.setFont("Helvetica", 7)
                c.drawCentredString(bx+bar_w/2, chart_bottom-4*mm, m)
            y = chart_bottom - 10*mm

    # footer
    c.setFillColor(GREEN_C); c.setFont("Helvetica", 7.5)
    c.drawString(15*mm, 12*mm, "Give A Bit · self-custody, no-KYC, sovereign · Confidential · For internal use")
    c.showPage(); c.save()

# ── Excel ────────────────────────────────────────────────────────────
def build_excel(title, site_name, report_id, generated, scope_lines, sections, path):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    wb = Workbook(); ws = wb.active; ws.title = "Report"
    ws.sheet_view.showGridLines = False
    ws.column_dimensions['A'].width = 40; ws.column_dimensions['B'].width = 24
    FILL = PatternFill("solid", fgColor="0E1B2A"); FILL2 = PatternFill("solid", fgColor="B8893A")
    ws["A1"] = title; ws["A1"].font = Font(bold=True, size=14, color="0E1B2A")
    ws["A2"] = f"Report ID: {report_id}"; ws["A2"].font = Font(italic=True, color="6B7280")
    ws["A3"] = f"Generated: {generated}"; ws["A3"].font = Font(italic=True, color="6B7280")
    r = 5
    ws.cell(row=r,column=1,value="REPORT SCOPE").fill=FILL; ws.cell(row=r,column=1).font=Font(bold=True,color="FFFFFF"); r+=1
    for ln in scope_lines:
        ws.cell(row=r,column=1,value=ln); ws.cell(row=r,column=1).font=Font(size=9); r+=1
    r += 1
    for t2, rows in sections:
        ws.cell(row=r,column=1,value=t2).fill=FILL2; ws.cell(row=r,column=1).font=Font(bold=True,color="FFFFFF"); r+=1
        for row in rows:
            if isinstance(row,(list,tuple)):
                ws.cell(row=r,column=1,value=row[0])
                if len(row)>1: ws.cell(row=r,column=2,value=row[1])
            else:
                ws.cell(row=r,column=1,value=row)
            r+=1
        r+=1
    wb.save(path)

# ── Word ─────────────────────────────────────────────────────────────
def build_word(title, site_name, report_id, generated, scope_lines, sections, path):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    doc = Document()
    try: doc.add_picture(LOGO, width=Inches(1.8))
    except Exception: pass
    doc.add_heading(title, 0)
    p=doc.add_paragraph(); p.add_run(f"Report ID: {report_id}   ·   Generated: {generated}").italic=True
    doc.add_heading("Report Scope", level=1)
    for ln in scope_lines: doc.add_paragraph(ln, style="List Bullet")
    for t2, rows in sections:
        doc.add_heading(t2, level=1)
        for row in rows:
            if isinstance(row,(list,tuple)):
                p=doc.add_paragraph(); p.add_run(f"{row[0]}: ").bold=True; p.add_run(str(row[1]))
            else:
                doc.add_paragraph(str(row), style="List Bullet")
    doc.save(path)

def render(site=None, outdir="/root/giveabit-reports"):
    sys.path.insert(0, "/root/hq")
    from report_insights import compute_insights, all_sites_comparison
    sht = build('sheets','v4',credentials=get_creds())
    os.makedirs(outdir, exist_ok=True)
    now = datetime.datetime.now()
    stamp = now.strftime("%Y%m%d-%H%M")
    generated = now.strftime("%Y-%m-%d %H:%M")

    if site:
        info = SITE_INFO[site]
        d = gather_site_data(sht, site)
        # Build insights for this site from its stamps
        site_stamps = []
        ledger = fetch_sheet(sht, MASTER, 'Ledger!A2:J1000')
        for r in ledger:
            if not r or not r[0]: continue
            cl = (r[7] if len(r)>7 else "").lower(); fn = (r[2] if len(r)>2 else "").lower()
            if site in cl or site in fn:
                site_stamps.append({"id": r[0], "status": r[3] if len(r)>3 else "", "client": r[7] if len(r)>7 else "", "created_at": r[4] if len(r)>4 else ""})
        pays = fetch_sheet(sht, MASTER, 'Payments!A2:L1000')
        site_pays = [r for r in pays if r and r[0] and str(r[0])[:4].isdigit() and (len(r)>1 and (r[1] or "").lower()==site)]
        ins = compute_insights(site_stamps, site_pays, base_price=d.get("base_price", 5))
        title = f"{site.title()} — Accounting Report"
        report_id = gen_report_id("accounting", site)
        scope_lines = build_scope_text()
        sections = [
            ("Executive Summary", [ln for ln in ins["exec_summary_lines"]]),
            ("Site", [("Project", site.title()), ("Domain", info["domain"]), ("Description", info["desc"])]),
            ("Key Figures", [
                ("Total .ots stamps", ins["stamps_total"]),
                ("Confirmed stamps", ins["stamps_confirmed"]),
                ("Payments recorded", ins["payments_count"]),
                ("Sats received (IN)", f"{ins['total_sats_in']:,}"),
            ]),
        ]
        if ins["anomalies"]:
            sections.append(("Flags & Anomalies", ins["anomalies"]))
        fname = f"giveabit-{site}-report-{stamp}"
    else:
        d = gather_master_data(sht)
        # Build insights for consolidated
        ledger = fetch_sheet(sht, MASTER, 'Ledger!A2:J1000')
        all_stamps = [{"id": r[0], "status": r[3] if len(r)>3 else "", "client": r[7] if len(r)>7 else "", "created_at": r[4] if len(r)>4 else ""} for r in ledger if r and r[0]]
        pays = fetch_sheet(sht, MASTER, 'Payments!A2:L1000')
        ins = compute_insights(all_stamps, pays, base_price=d["base_price"])
        title = "Give A Bit — Consolidated Accounting Report"
        report_id = gen_report_id("consolidated")
        scope_lines = build_scope_text()
        client_rows = sorted(ins["client_breakdown"].items(), key=lambda x:-x[1])
        sections = [
            ("Executive Summary", [ln for ln in ins["exec_summary_lines"]]),
            ("Key Figures", [
                ("Total .ots stamps", ins["stamps_total"]),
                ("Confirmed stamps", ins["stamps_confirmed"]),
                ("Pending stamps", ins["stamps_pending"]),
                ("Failed stamps", ins["stamps_failed"]),
                ("Distinct clients", ins["distinct_clients"]),
                ("Payments recorded", ins["payments_count"]),
                ("Total sats received (IN)", f"{ins['total_sats_in']:,}"),
            ]),
            ("Stamps by Client", client_rows),
            ("Pricing (global rate control)", [
                ("Base price/stamp (sats)", d["base_price"]),
                ("Rate adjustment (%)", d["rate_adj"]),
                ("Effective price/stamp (sats)", d["eff_price"]),
                ("BTC price (USD, live)", f"${d['btc_usd']:,.0f}"),
                ("Effective price in USD", f"${d['eff_price']*d['btc_usd']/100000000:.4f}"),
            ]),
        ]
        if ins["anomalies"]:
            sections.append(("Flags & Anomalies", ins["anomalies"]))
        # All-sites comparison table
        cmp_rows = []
        for sname in SITE_INFO:
            sd = gather_site_data(sht, sname)
            cmp_rows.append([sname.title(), sd["stamps_total"], sd["stamps_confirmed"], sd["payments"], f"{sd['sats_in']:,}"])
        sections.append(("All Sites Comparison", [["Site", "Stamps", "Confirmed", "Payments", "Sats IN"]] + cmp_rows))
        # Trend chart data (monthly stamps)
        trend = {"monthly": ins.get("monthly", {})}
        fname = f"giveabit-consolidated-report-{stamp}"

    pdf = f"{outdir}/{fname}.pdf"
    xls = f"{outdir}/{fname}.xlsx"
    doc = f"{outdir}/{fname}.docx"
    build_pdf(title, site.title() if site else "Consolidated — all projects", report_id, generated, scope_lines, sections, pdf, trend=locals().get("trend"))
    build_excel(title, site.title() if site else "Consolidated", report_id, generated, scope_lines, sections, xls)
    build_word(title, site.title() if site else "Consolidated", report_id, generated, scope_lines, sections, doc)
    return {"report_id": report_id, "pdf": pdf, "xlsx": xls, "docx": doc, "site": site, "title": title}

if __name__ == "__main__":
    import argparse
    ap = argparse.ArgumentParser()
    ap.add_argument("--site", help="generate per-site report (motopass, satohash, katoa, tadbuy, sherpacarta, stranded)")
    ap.add_argument("--all-sites", action="store_true", help="generate all per-site reports")
    ap.add_argument("--master", action="store_true", help="generate consolidated master report")
    args = ap.parse_args()
    results = []
    if args.master or not (args.site or args.all_sites):
        results.append(render())
    if args.all_sites:
        for s in SITE_INFO: results.append(render(site=s))
    elif args.site:
        results.append(render(site=args.site))
    for r in results:
        print(f"{r['site'] or 'CONSOLIDATED':14} {r['report_id']}")
        print(f"   {r['pdf']}")
        print(f"   {r['xlsx']}")
        print(f"   {r['docx']}")
